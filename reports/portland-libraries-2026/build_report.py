from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Resolve relative to this file so the report builds from any checkout,
# not only the machine it was first written on.
ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "report-source.md"
OUTPUT = ROOT / "Portland_Libraries_Deep_Dive_2026.docx"

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "26766A"
GOLD = "B27A22"
INK = "1F2933"
MUTED = "66727D"
LIGHT = "F4F6F9"
PALE_BLUE = "EAF2F7"
PALE_TEAL = "EAF4F1"
WHITE = "FFFFFF"
BORDER = "CBD5DD"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, side="bottom", color=BORDER, size="8", space="3") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    borders.append(border)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, *, bold=False, italic=False, size=None, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        r_pr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(int(size * 2)))
        r_pr.append(sz_cs)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text: str, *, default_size=None, default_color=INK) -> None:
    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, color=default_color, italic=True)
        else:
            m = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            add_hyperlink(paragraph, m.group(1), m.group(2), size=default_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size, color=default_color)


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)


def new_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "279")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Table Citation" not in styles:
        styles.add_style("Table Citation", 1)
    citation = styles["Table Citation"]
    citation.font.name = "Calibri"
    citation._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    citation._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    citation.font.size = Pt(8.5)
    citation.font.color.rgb = rgb(MUTED)
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    citation.paragraph_format.line_spacing = 1.0


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    run = hp.add_run("PORTLAND PUBLIC KNOWLEDGE COMMONS")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(hp, color=BORDER, size="6", space="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = fp.add_run("RESEARCH & STRATEGY  ·  AUGUST 2026   |   ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_number(fp)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(118)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CIVIC STRATEGY REPORT")
    set_run_font(r, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Portland Public\nKnowledge Commons")
    set_run_font(r, size=31, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("A deep dive into the past, present, and possible future\nof Multnomah County Library")
    set_run_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("1864–2026 history  ·  2011–2026 transformation  ·  2040 horizon")
    set_run_font(r, size=9.5, color=TEAL, bold=True)

    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_TEAL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("THE CENTRAL QUESTION")
    set_run_font(r, size=8.5, color=TEAL, bold=True)
    p.add_run("\n")
    r = p.add_run("How can Portland turn a rebuilt library network into the world’s most effective public infrastructure for human agency and democratic belonging?")
    set_run_font(r, size=11.5, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared August 2026  |  Multnomah County, Oregon")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    h = doc.add_heading("Reader’s map", level=1)
    h.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    add_inline(p, "This report separates the historical record, the current diagnosis, and the proposed future. The future model and numeric targets are recommendations—not claims about current performance.")

    rows = [
        ("Executive answer", "The thesis, the recent transformation, and the strategic choice."),
        ("1 · Public purpose", "Six recurring purposes and the tension inside public librarianship."),
        ("2 · Portland history", "From subscription privilege to a countywide public network."),
        ("3 · Last fifteen years", "Permanent funding, hybrid use, crisis, and reconstruction."),
        ("4 · The gap", "Where Portland is strong and what separates it from world leadership."),
        ("5 · Global lessons", "Practice benchmarks from Helsinki, Aarhus, Singapore, Toronto, Vancouver, Christchurch, Chicago, and San Francisco."),
        ("6 · Future model", "The Portland Public Knowledge Commons: three layers and five protected purposes."),
        ("7 · World-leading standard", "A 2040 scorecard with equity floors and anti-gaming rules."),
        ("8 · Roadmap", "A staged path from stabilization to independently verified outcomes."),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2340, 7020])
    for i, (label, detail) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_shading(cell, LIGHT if i % 2 == 0 else WHITE)
        p1 = table.cell(i, 0).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_after = Pt(1)
        r = p1.add_run(label)
        set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
        p2 = table.cell(i, 1).paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_after = Pt(1)
        r = p2.add_run(detail)
        set_run_font(r, size=9.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SCOPE")
    set_run_font(r, size=8.5, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "“Portland libraries” means Multnomah County Library: a countywide system serving Portland, Fairview, Gresham, Troutdale, and unincorporated communities. The Library District is a separate taxing district, while the County Board governs it and MCL operates as a county department.", default_size=9.5, default_color=MUTED)
    doc.add_page_break()


def add_thesis_callout(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("BOTTOM LINE\n")
    set_run_font(r, size=8.5, color=BLUE, bold=True)
    r = p.add_run("Portland has much of the hardware of a world-leading system. The gap is the operating model: universal resident reach, branch-by-branch equity, staff safety, specialist partnerships, and credible evidence of learning and civic outcomes.")
    set_run_font(r, size=11, color=NAVY, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.line_spacing = 0.2
    r = spacer.add_run("\u00a0")
    set_run_font(r, size=2, color=WHITE)


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_data_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    if cols == 6:
        widths = [1980, 1170, 1170, 1170, 1170, 2700]
        font_size = 7.5
    elif cols == 4:
        widths = [1350, 2370, 2520, 3120]
        font_size = 7.9
    elif cols == 3:
        widths = [1620, 2880, 4860]
        font_size = 8.2
    elif cols == 2:
        widths = [2700, 6660]
        font_size = 8.7
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
        font_size = 8
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            set_cell_shading(cell, NAVY if i == 0 else (LIGHT if i % 2 == 0 else WHITE))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, default_size=font_size, default_color=WHITE if i == 0 else INK)
            for run in p.runs:
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = rgb(WHITE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_body_paragraph(doc: Document, text: str, *, table_citation=False) -> None:
    p = doc.add_paragraph(style="Table Citation" if table_citation else "Normal")
    add_inline(p, text, default_size=8.5 if table_citation else None, default_color=MUTED if table_citation else INK)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    add_contents(doc)

    # Skip source-title furniture and begin with Executive answer.
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Executive answer")
    i = start
    bullet_num = None
    decimal_num = None
    in_list_type = None
    previous_was_table = False
    major_sections = 0

    while i < len(lines):
        raw = lines[i].rstrip()
        text = raw.strip()
        if not text:
            in_list_type = None
            previous_was_table = False
            i += 1
            continue

        if text.startswith("|"):
            rows, next_i = parse_table(lines, i)
            if len(rows[0]) == 6:
                doc.add_page_break()
            add_data_table(doc, rows)
            previous_was_table = True
            i = next_i
            continue

        if text.startswith("## "):
            title = text[3:]
            if major_sections > 0:
                doc.add_page_break()
            h = doc.add_heading(title, level=1)
            h.paragraph_format.space_before = Pt(0)
            if title == "Executive answer":
                add_thesis_callout(doc)
            major_sections += 1
            previous_was_table = False
            i += 1
            continue

        if text.startswith("### "):
            doc.add_heading(text[4:], level=2)
            previous_was_table = False
            i += 1
            continue

        if text.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            set_table_geometry(table, [9360])
            cell = table.cell(0, 0)
            set_cell_shading(cell, PALE_TEAL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, text[2:], default_size=11.5, default_color=NAVY)
            for run in p.runs:
                run.bold = True
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(4)
            spacer.paragraph_format.line_spacing = 0.2
            r = spacer.add_run("\u00a0")
            set_run_font(r, size=2, color=WHITE)
            i += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", text)
        numbered = re.match(r"^(\d+)\.\s+(.+)$", text)
        if bullet:
            if in_list_type != "bullet":
                bullet_num = new_numbering(doc, "bullet")
                in_list_type = "bullet"
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            apply_numbering(p, bullet_num)
            add_inline(p, bullet.group(1))
            previous_was_table = False
            i += 1
            continue
        if numbered:
            if in_list_type != "decimal":
                decimal_num = new_numbering(doc, "decimal")
                in_list_type = "decimal"
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            apply_numbering(p, decimal_num)
            add_inline(p, numbered.group(2))
            previous_was_table = False
            i += 1
            continue

        is_citation = previous_was_table and (text.startswith("*") or text.startswith("Source:"))
        add_body_paragraph(doc, text, table_citation=is_citation)
        previous_was_table = False
        in_list_type = None
        i += 1

    props = doc.core_properties
    props.title = "Portland Public Knowledge Commons"
    props.subject = "A deep dive into the history, transformation, and possible future of Multnomah County Library"
    props.author = "OpenAI"
    props.keywords = "Portland, Multnomah County Library, public libraries, civic infrastructure, library strategy"
    props.comments = "Research and strategy report prepared August 2026"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
